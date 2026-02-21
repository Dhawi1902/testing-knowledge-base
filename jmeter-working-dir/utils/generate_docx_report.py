#!/usr/bin/env python3
"""
MAYA Performance Test DOCX Report Generator

Generates professional DOCX reports from JMeter test results.
Filters out student IDs and embedded resources automatically.

Usage:
  # Single run
  python utils/generate_docx_report.py results/jmeter-report/20260204/20260204_10

  # All runs in a date folder
  python utils/generate_docx_report.py results/jmeter-report/20260204

  # Custom output directory
  python utils/generate_docx_report.py results/jmeter-report/20260204 --output reports/docx/
"""

import json
import os
import re
import sys
import argparse
import tempfile
from datetime import datetime
from pathlib import Path

import pandas as pd
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.ticker as ticker
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# --- Constants ---
STUDENT_ID_PATTERN = re.compile(r'^[A-Z]{2,4}\d{4,7}$')
EMBEDDED_SUFFIX_PATTERN = re.compile(r'-\d+$')
THEME_COLOR = RGBColor(0x2E, 0x74, 0xB5)
THEME_COLOR_HEX = '2E74B5'
HEADER_BG_HEX = '2E74B5'
ALT_ROW_BG_HEX = 'D6E4F0'
ERROR_COLOR = RGBColor(0xC0, 0x00, 0x00)
SUCCESS_COLOR = RGBColor(0x00, 0x80, 0x00)


# --- Filtering ---

def is_student_id(label):
    """Check if a label is a student ID (e.g., EBR003819)."""
    return bool(STUDENT_ID_PATTERN.match(label.strip()))


def is_embedded_resource(label):
    """Check if a label is an embedded resource (e.g., 00 - GET - Login_Page-3)."""
    return bool(EMBEDDED_SUFFIX_PATTERN.search(label.strip()))


def filter_transactions(stats):
    """Filter statistics to keep only meaningful parent transactions."""
    filtered = {}
    for label, data in stats.items():
        clean_label = label.strip()
        if is_student_id(clean_label):
            continue
        if is_embedded_resource(clean_label):
            continue
        data['transaction'] = clean_label
        filtered[clean_label] = data
    return filtered


def sort_transactions(transactions):
    """Sort transactions: numbered first (by number), then named, then Total last."""
    def sort_key(label):
        if label == 'Total':
            return (2, 0, label)
        if label == 'Student Enrolment':
            return (1, 0, label)
        match = re.match(r'^(\d+)', label)
        if match:
            return (0, int(match.group(1)), label)
        return (0, 999, label)

    return sorted(transactions.keys(), key=sort_key)


# --- Data Loading ---

def load_from_statistics_json(stats_path):
    """Load pre-aggregated statistics from JMeter HTML report."""
    with open(stats_path, 'r', encoding='utf-8') as f:
        return json.load(f)


def load_from_jtl(jtl_path):
    """Compute statistics from raw JTL CSV file (fallback)."""
    print(f"  Parsing JTL file (this may take a while for large files)...")
    df = pd.read_csv(jtl_path, low_memory=False)
    stats = {}
    for label, group in df.groupby('label'):
        elapsed = group['elapsed']
        error_count = int((group['success'] == False).sum())  # noqa: E712
        sample_count = len(group)
        duration_s = max((group['timeStamp'].max() - group['timeStamp'].min()) / 1000.0, 1)
        stats[label] = {
            'transaction': label,
            'sampleCount': sample_count,
            'errorCount': error_count,
            'errorPct': round(error_count / sample_count * 100, 2) if sample_count > 0 else 0,
            'meanResTime': round(float(elapsed.mean()), 2),
            'medianResTime': round(float(elapsed.median()), 2),
            'minResTime': float(elapsed.min()),
            'maxResTime': float(elapsed.max()),
            'pct1ResTime': round(float(elapsed.quantile(0.90)), 2),
            'pct2ResTime': round(float(elapsed.quantile(0.95)), 2),
            'pct3ResTime': round(float(elapsed.quantile(0.99)), 2),
            'throughput': round(sample_count / duration_s, 2),
            'receivedKBytesPerSec': 0,
            'sentKBytesPerSec': 0,
        }
    return stats


def extract_metadata(run_dir):
    """Extract test metadata from available sources."""
    run_dir = Path(run_dir)
    meta = {
        'run_id': run_dir.name,
        'date': '',
        'start_time': '',
        'end_time': '',
        'duration': '',
    }

    # Try to extract date from folder name (e.g., 20260204_10)
    match = re.match(r'(\d{8})_(\d+)', run_dir.name)
    if match:
        date_str = match.group(1)
        try:
            meta['date'] = datetime.strptime(date_str, '%Y%m%d').strftime('%d %B %Y')
        except ValueError:
            meta['date'] = date_str
        meta['run_number'] = match.group(2)

    # Try to get start/end from HTML report
    index_html = run_dir / 'report' / 'index.html'
    if index_html.exists():
        try:
            content = index_html.read_text(encoding='utf-8', errors='ignore')
            start_match = re.search(r'Start Time.*?<td>.*?"([^"]+)"', content, re.DOTALL)
            end_match = re.search(r'End Time.*?<td>.*?"([^"]+)"', content, re.DOTALL)
            if start_match:
                meta['start_time'] = start_match.group(1)
            if end_match:
                meta['end_time'] = end_match.group(1)
        except Exception:
            pass

    # Fallback: extract from JTL timestamps
    if not meta['start_time']:
        jtl_path = run_dir / 'results.jtl'
        if jtl_path.exists():
            try:
                df = pd.read_csv(jtl_path, nrows=5, usecols=['timeStamp'])
                start_ts = df['timeStamp'].min()
                meta['start_time'] = datetime.fromtimestamp(start_ts / 1000).strftime('%d/%m/%Y %I:%M %p')
            except Exception:
                pass

    return meta


# --- Chart Generation ---

def generate_response_time_chart(transactions, sorted_labels, output_path):
    """Generate response time bar chart."""
    # Exclude Total and Student Enrolment from chart
    chart_labels = [l for l in sorted_labels if l not in ('Total', 'Student Enrolment')]
    if not chart_labels:
        return False

    # Shorten labels for display
    short_labels = []
    for l in chart_labels:
        short = re.sub(r'^\d+ - (GET|POST) - ', '', l)
        if len(short) > 25:
            short = short[:22] + '...'
        short_labels.append(short)

    mean_times = [transactions[l]['meanResTime'] for l in chart_labels]
    p90_times = [transactions[l].get('pct1ResTime', 0) for l in chart_labels]
    p95_times = [transactions[l].get('pct2ResTime', 0) for l in chart_labels]

    fig, ax = plt.subplots(figsize=(10, max(5, len(chart_labels) * 0.4)))
    y_pos = range(len(chart_labels))
    bar_height = 0.25

    bars3 = ax.barh([y - bar_height for y in y_pos], p95_times, bar_height,
                    label='95th Percentile', color='#E8927C', alpha=0.85)
    bars2 = ax.barh(y_pos, p90_times, bar_height,
                    label='90th Percentile', color='#F4C542', alpha=0.85)
    bars1 = ax.barh([y + bar_height for y in y_pos], mean_times, bar_height,
                    label='Mean', color='#5B9BD5', alpha=0.85)

    ax.set_yticks(y_pos)
    ax.set_yticklabels(short_labels, fontsize=8)
    ax.set_xlabel('Response Time (ms)', fontsize=10)
    ax.set_title('Response Time by Transaction', fontsize=12, fontweight='bold')
    ax.legend(loc='lower right', fontsize=8)
    ax.xaxis.set_major_formatter(ticker.FuncFormatter(lambda x, _: f'{x:,.0f}'))
    ax.invert_yaxis()
    plt.tight_layout()
    fig.savefig(output_path, dpi=150, bbox_inches='tight')
    plt.close(fig)
    return True


def generate_error_rate_chart(transactions, sorted_labels, output_path):
    """Generate error rate bar chart."""
    chart_labels = [l for l in sorted_labels if l not in ('Total', 'Student Enrolment')]
    if not chart_labels:
        return False

    error_pcts = [transactions[l].get('errorPct', 0) for l in chart_labels]
    if max(error_pcts) == 0:
        return False

    short_labels = []
    for l in chart_labels:
        short = re.sub(r'^\d+ - (GET|POST) - ', '', l)
        if len(short) > 25:
            short = short[:22] + '...'
        short_labels.append(short)

    fig, ax = plt.subplots(figsize=(10, max(4, len(chart_labels) * 0.35)))
    colors = ['#C00000' if e > 5 else '#F4C542' if e > 0 else '#70AD47' for e in error_pcts]
    ax.barh(range(len(chart_labels)), error_pcts, color=colors, alpha=0.85)
    ax.set_yticks(range(len(chart_labels)))
    ax.set_yticklabels(short_labels, fontsize=8)
    ax.set_xlabel('Error Rate (%)', fontsize=10)
    ax.set_title('Error Rate by Transaction', fontsize=12, fontweight='bold')
    ax.invert_yaxis()
    plt.tight_layout()
    fig.savefig(output_path, dpi=150, bbox_inches='tight')
    plt.close(fig)
    return True


def generate_throughput_chart(transactions, sorted_labels, output_path):
    """Generate throughput bar chart."""
    chart_labels = [l for l in sorted_labels if l not in ('Total', 'Student Enrolment')]
    if not chart_labels:
        return False

    short_labels = []
    for l in chart_labels:
        short = re.sub(r'^\d+ - (GET|POST) - ', '', l)
        if len(short) > 25:
            short = short[:22] + '...'
        short_labels.append(short)

    throughputs = [transactions[l].get('throughput', 0) for l in chart_labels]

    fig, ax = plt.subplots(figsize=(10, max(4, len(chart_labels) * 0.35)))
    ax.barh(range(len(chart_labels)), throughputs, color='#5B9BD5', alpha=0.85)
    ax.set_yticks(range(len(chart_labels)))
    ax.set_yticklabels(short_labels, fontsize=8)
    ax.set_xlabel('Throughput (req/sec)', fontsize=10)
    ax.set_title('Throughput by Transaction', fontsize=12, fontweight='bold')
    ax.invert_yaxis()
    plt.tight_layout()
    fig.savefig(output_path, dpi=150, bbox_inches='tight')
    plt.close(fig)
    return True


# --- DOCX Building ---

def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, color=None, size=Pt(9), alignment=None):
    """Set text in a table cell with formatting."""
    cell.text = ''
    paragraph = cell.paragraphs[0]
    if alignment:
        paragraph.alignment = alignment
    run = paragraph.add_run(str(text))
    run.font.size = size
    run.font.name = 'Calibri'
    if bold:
        run.font.bold = True
    if color:
        run.font.color.rgb = color


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a professionally styled table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, HEADER_BG_HEX)
        set_cell_text(cell, header, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=Pt(9))

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            if row_idx % 2 == 1:
                set_cell_shading(cell, ALT_ROW_BG_HEX)

            # Format numbers
            if isinstance(value, float):
                if value > 1000:
                    text = f'{value:,.0f}'
                else:
                    text = f'{value:.2f}'
            elif isinstance(value, int):
                text = f'{value:,}'
            else:
                text = str(value)

            is_error_col = headers[col_idx] in ('Error %', 'Errors')
            color = None
            if is_error_col and isinstance(value, (int, float)) and value > 0:
                color = ERROR_COLOR

            alignment = WD_ALIGN_PARAGRAPH.RIGHT if isinstance(value, (int, float)) else None
            set_cell_text(cell, text, size=Pt(8), color=color, alignment=alignment)

    # Apply column widths
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = width

    return table


def build_docx(meta, transactions, sorted_labels, chart_paths, output_path):
    """Build the complete DOCX report."""
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # --- Default font ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    # =====================
    # COVER PAGE
    # =====================
    for _ in range(6):
        doc.add_paragraph('')

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('MAYA Portal')
    run.font.size = Pt(28)
    run.font.color.rgb = THEME_COLOR
    run.font.bold = True
    run.font.name = 'Calibri'

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Performance Test Report')
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    run.font.name = 'Calibri'

    doc.add_paragraph('')

    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    details = [
        f"Run ID: {meta.get('run_id', 'N/A')}",
        f"Date: {meta.get('date', 'N/A')}",
        f"Environment: Production (maya-cloud.um.edu.my)",
    ]
    for detail in details:
        run = info_para.add_run(detail + '\n')
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        run.font.name = 'Calibri'

    doc.add_page_break()

    # =====================
    # TEST CONFIGURATION
    # =====================
    h = doc.add_heading('1. Test Configuration', level=1)
    h.runs[0].font.color.rgb = THEME_COLOR

    config_headers = ['Parameter', 'Value']
    config_rows = [
        ('Run ID', meta.get('run_id', 'N/A')),
        ('Date', meta.get('date', 'N/A')),
        ('Start Time', meta.get('start_time', 'N/A')),
        ('End Time', meta.get('end_time', 'N/A')),
        ('Target System', 'MAYA Portal (SITS:Vision)'),
        ('Environment', 'Production'),
        ('URL', 'https://maya-cloud.um.edu.my'),
    ]
    add_styled_table(doc, config_headers, config_rows,
                     col_widths=[Cm(5), Cm(12)])

    doc.add_paragraph('')

    # =====================
    # EXECUTIVE SUMMARY
    # =====================
    h = doc.add_heading('2. Executive Summary', level=1)
    h.runs[0].font.color.rgb = THEME_COLOR

    total = transactions.get('Total', {})
    student_enrolment = transactions.get('Student Enrolment', {})

    summary_headers = ['Metric', 'Value']
    summary_rows = [
        ('Total Samples', total.get('sampleCount', 'N/A')),
        ('Total Errors', total.get('errorCount', 'N/A')),
        ('Error Rate (%)', f"{total.get('errorPct', 0):.2f}%"),
        ('Mean Response Time (ms)', f"{total.get('meanResTime', 0):,.0f}"),
        ('90th Percentile (ms)', f"{total.get('pct1ResTime', 0):,.0f}"),
        ('95th Percentile (ms)', f"{total.get('pct2ResTime', 0):,.0f}"),
        ('Throughput (req/sec)', f"{total.get('throughput', 0):,.2f}"),
    ]
    if student_enrolment:
        summary_rows.insert(0, ('Workflow', 'Student Enrolment'))
        summary_rows.append(('Avg Workflow Duration (ms)',
                            f"{student_enrolment.get('meanResTime', 0):,.0f}"))

    add_styled_table(doc, summary_headers, summary_rows,
                     col_widths=[Cm(7), Cm(10)])

    # Verdict
    error_pct = total.get('errorPct', 0)
    doc.add_paragraph('')
    verdict_para = doc.add_paragraph()
    if error_pct <= 2:
        run = verdict_para.add_run('PASS')
        run.font.color.rgb = SUCCESS_COLOR
        run.font.bold = True
        run.font.size = Pt(14)
        verdict_para.add_run(' - Error rate within acceptable threshold (< 2%)')
    elif error_pct <= 5:
        run = verdict_para.add_run('WARNING')
        run.font.color.rgb = RGBColor(0xFF, 0xA5, 0x00)
        run.font.bold = True
        run.font.size = Pt(14)
        verdict_para.add_run(f' - Error rate at {error_pct:.1f}% (threshold: < 2%)')
    else:
        run = verdict_para.add_run('FAIL')
        run.font.color.rgb = ERROR_COLOR
        run.font.bold = True
        run.font.size = Pt(14)
        verdict_para.add_run(f' - Error rate at {error_pct:.1f}% exceeds threshold (> 5%)')

    doc.add_page_break()

    # =====================
    # RESPONSE TIME SUMMARY
    # =====================
    h = doc.add_heading('3. Response Time Summary', level=1)
    h.runs[0].font.color.rgb = THEME_COLOR

    rt_headers = ['Transaction', 'Samples', 'Mean (ms)', 'Median (ms)',
                  'P90 (ms)', 'P95 (ms)', 'P99 (ms)', 'Min (ms)', 'Max (ms)']

    rt_rows = []
    for label in sorted_labels:
        t = transactions[label]
        rt_rows.append((
            label,
            t.get('sampleCount', 0),
            t.get('meanResTime', 0),
            t.get('medianResTime', 0),
            t.get('pct1ResTime', 0),
            t.get('pct2ResTime', 0),
            t.get('pct3ResTime', 0),
            t.get('minResTime', 0),
            t.get('maxResTime', 0),
        ))

    add_styled_table(doc, rt_headers, rt_rows,
                     col_widths=[Cm(4.5), Cm(1.5), Cm(1.8), Cm(1.8),
                                 Cm(1.5), Cm(1.5), Cm(1.5), Cm(1.5), Cm(1.5)])

    doc.add_paragraph('')

    # Response time chart
    if 'response_time' in chart_paths and chart_paths['response_time']:
        doc.add_picture(chart_paths['response_time'], width=Cm(17))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # =====================
    # THROUGHPUT
    # =====================
    h = doc.add_heading('4. Throughput', level=1)
    h.runs[0].font.color.rgb = THEME_COLOR

    tp_headers = ['Transaction', 'Samples', 'Throughput (req/s)',
                  'Received (KB/s)', 'Sent (KB/s)']

    tp_rows = []
    for label in sorted_labels:
        t = transactions[label]
        tp_rows.append((
            label,
            t.get('sampleCount', 0),
            round(t.get('throughput', 0), 2),
            round(t.get('receivedKBytesPerSec', 0), 2),
            round(t.get('sentKBytesPerSec', 0), 2),
        ))

    add_styled_table(doc, tp_headers, tp_rows,
                     col_widths=[Cm(5), Cm(2.5), Cm(3), Cm(3), Cm(3)])

    doc.add_paragraph('')

    if 'throughput' in chart_paths and chart_paths['throughput']:
        doc.add_picture(chart_paths['throughput'], width=Cm(17))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # =====================
    # ERROR ANALYSIS
    # =====================
    h = doc.add_heading('5. Error Analysis', level=1)
    h.runs[0].font.color.rgb = THEME_COLOR

    err_headers = ['Transaction', 'Samples', 'Errors', 'Error %']
    err_rows = []
    for label in sorted_labels:
        t = transactions[label]
        err_rows.append((
            label,
            t.get('sampleCount', 0),
            t.get('errorCount', 0),
            round(t.get('errorPct', 0), 2),
        ))

    add_styled_table(doc, err_headers, err_rows,
                     col_widths=[Cm(7), Cm(3), Cm(3), Cm(3)])

    doc.add_paragraph('')

    if 'error_rate' in chart_paths and chart_paths['error_rate']:
        doc.add_picture(chart_paths['error_rate'], width=Cm(17))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # =====================
    # FOOTER
    # =====================
    doc.add_paragraph('')
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f'Report generated on {datetime.now().strftime("%d %B %Y %H:%M:%S")}')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(output_path)
    return output_path


# --- Main Pipeline ---

def process_single_run(run_dir, output_dir=None):
    """Process a single test run and generate DOCX report."""
    run_dir = Path(run_dir)
    run_name = run_dir.name

    if output_dir:
        output_dir = Path(output_dir)
    else:
        output_dir = run_dir

    output_dir.mkdir(parents=True, exist_ok=True)

    print(f"\nProcessing: {run_name}")

    # 1. Load statistics
    stats_path = run_dir / 'report' / 'statistics.json'
    jtl_path = run_dir / 'results.jtl'

    if stats_path.exists():
        print(f"  Loading statistics.json...")
        raw_stats = load_from_statistics_json(stats_path)
    elif jtl_path.exists():
        file_size_mb = jtl_path.stat().st_size / (1024 * 1024)
        if file_size_mb < 1:
            print(f"  WARNING: JTL file is very small ({file_size_mb:.1f} MB) - possibly incomplete run")
        print(f"  No statistics.json found. Parsing JTL ({file_size_mb:.0f} MB)...")
        raw_stats = load_from_jtl(jtl_path)
    else:
        print(f"  SKIPPED: No statistics.json or results.jtl found")
        return None

    # 2. Filter transactions
    transactions = filter_transactions(raw_stats)
    if not transactions:
        print(f"  SKIPPED: No valid transactions found after filtering")
        return None

    sorted_labels = sort_transactions(transactions)
    print(f"  Found {len(sorted_labels)} transactions (filtered from {len(raw_stats)} total entries)")

    # 3. Extract metadata
    meta = extract_metadata(run_dir)

    # 4. Generate charts
    chart_paths = {}
    with tempfile.TemporaryDirectory() as tmpdir:
        rt_chart = os.path.join(tmpdir, 'response_times.png')
        if generate_response_time_chart(transactions, sorted_labels, rt_chart):
            chart_paths['response_time'] = rt_chart

        err_chart = os.path.join(tmpdir, 'error_rate.png')
        if generate_error_rate_chart(transactions, sorted_labels, err_chart):
            chart_paths['error_rate'] = err_chart

        tp_chart = os.path.join(tmpdir, 'throughput.png')
        if generate_throughput_chart(transactions, sorted_labels, tp_chart):
            chart_paths['throughput'] = tp_chart

        # 5. Build DOCX
        output_path = output_dir / f'{run_name}_report.docx'
        build_docx(meta, transactions, sorted_labels, chart_paths, str(output_path))

    print(f"  Report saved: {output_path}")
    return output_path


def process_batch(parent_dir, output_dir=None):
    """Process all run folders in a parent directory."""
    parent_dir = Path(parent_dir)
    reports = []

    # Find run folders (pattern: YYYYMMDD_N)
    run_dirs = sorted([
        d for d in parent_dir.iterdir()
        if d.is_dir() and re.match(r'\d{8}_\d+', d.name)
    ], key=lambda d: (d.name.split('_')[0], int(d.name.split('_')[1])))

    if not run_dirs:
        # Maybe it's a single run dir
        if (parent_dir / 'results.jtl').exists() or (parent_dir / 'report' / 'statistics.json').exists():
            result = process_single_run(parent_dir, output_dir)
            if result:
                reports.append(result)
            return reports
        print(f"No run folders found in {parent_dir}")
        return reports

    print(f"Found {len(run_dirs)} run folders in {parent_dir.name}")

    for run_dir in run_dirs:
        try:
            result = process_single_run(run_dir, output_dir)
            if result:
                reports.append(result)
        except Exception as e:
            print(f"  ERROR processing {run_dir.name}: {e}")

    return reports


def main():
    parser = argparse.ArgumentParser(
        description='Generate DOCX performance test reports from JMeter results')
    parser.add_argument('input_path',
                        help='Path to a single run folder or parent directory containing multiple runs')
    parser.add_argument('--output', '-o',
                        help='Output directory for generated reports (default: same as input)')
    args = parser.parse_args()

    input_path = Path(args.input_path)
    if not input_path.exists():
        print(f"Error: {input_path} does not exist")
        sys.exit(1)

    output_dir = Path(args.output) if args.output else None

    reports = process_batch(input_path, output_dir)

    print(f"\n{'='*50}")
    print(f"Generated {len(reports)} report(s)")
    for r in reports:
        print(f"  - {r}")


if __name__ == '__main__':
    main()
